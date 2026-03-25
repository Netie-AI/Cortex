import os
try:
    from docx import Document
except ImportError:
    pass

def extract_docx(filepath, outpath):
    doc = Document(filepath)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            fullText.append(" | ".join(row_data))
    with open(outpath, 'w', encoding='utf-8') as f:
        f.write('\n'.join(fullText))

extract_docx('Netie_Cortex_v1_Architecture.docx', 'v1_arch.txt')
extract_docx('Netie_Cortex_v2_Strategic_Architecture.docx', 'v2_arch.txt')
